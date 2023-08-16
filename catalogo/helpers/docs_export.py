from docx import Document

def generate_word_doc(file_name):
    doc = Document()
    doc.add_heading('Ejemplo de Documento de Word', 0)
    doc.add_paragraph('Este es un párrafo de ejemplo.')
    doc.save(file_name)

generate_word_doc("ejemplo.docx")